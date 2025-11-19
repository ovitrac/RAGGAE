"""
RAGGAE/io/textloaders.py

Text document loaders for multiple formats (TXT, MD, DOCX, ODT).
Provides unified interface for extracting text blocks from various
document types with basic provenance tracking.

This module provides:
- TextBlock: Generic text block dataclass
- Format-specific loaders: load_txt, load_md, load_docx, load_odt
- Dispatcher: load_blocks_any for automatic format detection

Author: Dr. Olivier Vitrac, PhD, HDR
Email: olivier.vitrac@adservio.com
Organization: Adservio
Date: October 31, 2025
License: MIT

Examples
--------
>>> from io.textloaders import load_blocks_any
>>> blocks = load_blocks_any("report.docx", min_chars=20)
>>> print(f"Loaded {len(blocks)} blocks")
"""
# RAGGAE/io/textloaders.py
from __future__ import annotations
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple

@dataclass
class TextBlock:
    """Generic text block with provenance similar to PDFBlock."""
    text: str
    page: int  # always 1 for non-PDF
    block: int
    bbox: Tuple[float, float, float, float] = (0,0,0,0)

    def as_metadata(self) -> Dict:
        return {"page": self.page, "block": self.block, "bbox": self.bbox}

# ---- TXT / MD ----
def load_txt(path: str | Path, min_chars: int = 20) -> List[TextBlock]:
    p = Path(path)
    blocks: List[TextBlock] = []
    text = p.read_text(encoding="utf-8", errors="ignore")
    for i, para in enumerate([t.strip() for t in text.splitlines()]):
        if len(para) >= min_chars:
            blocks.append(TextBlock(text=para, page=1, block=i))
    return blocks

def load_md(path: str | Path, min_chars: int = 20) -> List[TextBlock]:
    # Same as txt (Markdown kept raw; no rendering)
    return load_txt(path, min_chars=min_chars)

# ---- DOCX ----
def load_docx(path: str | Path, min_chars: int = 20) -> List[TextBlock]:
    try:
        import docx  # python-docx
    except Exception as e:
        raise ImportError("python-docx not installed. `mamba install -c conda-forge python-docx`") from e
    p = Path(path)
    doc = docx.Document(str(p))
    blocks: List[TextBlock] = []
    bi = 0
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if len(t) >= min_chars:
            blocks.append(TextBlock(text=t, page=1, block=bi)); bi += 1
    # tables -> rows joined by ' | '
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            t = " | ".join([c for c in cells if c])
            if len(t) >= min_chars:
                blocks.append(TextBlock(text=t, page=1, block=bi)); bi += 1
    return blocks

# ---- ODT ----
def load_odt(path: str | Path, min_chars: int = 20) -> List[TextBlock]:
    try:
        from odf import text, teletype
        from odf.opendocument import load
    except Exception as e:
        raise ImportError("odfpy not installed. `mamba install -c conda-forge odfpy`") from e
    doc = load(str(path))
    paras = doc.getElementsByType(text.P)
    blocks: List[TextBlock] = []
    bi = 0
    for p in paras:
        t = (teletype.extractText(p) or "").strip()
        if len(t) >= min_chars:
            blocks.append(TextBlock(text=t, page=1, block=bi)); bi += 1
    return blocks

# ---- Dispatcher ----
def load_blocks_any(path: str | Path, min_chars: int = 20) -> List[TextBlock]:
    p = Path(path)
    ext = p.suffix.lower().lstrip(".")
    if ext == "txt":  return load_txt(p, min_chars=min_chars)
    if ext == "md":   return load_md(p, min_chars=min_chars)
    if ext == "docx": return load_docx(p, min_chars=min_chars)
    if ext == "odt":  return load_odt(p, min_chars=min_chars)
    raise ValueError(f"Unsupported extension for generic loader: .{ext}")
